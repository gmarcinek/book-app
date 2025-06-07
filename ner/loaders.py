"""
NER Text Loaders - Multi-format document loading
Supports: txt, md, pdf, docx, rtf with memory management
"""

import os
from pathlib import Path
from typing import Optional, Dict, Any
from dataclasses import dataclass

# Core imports
from .utils import (
    log_memory_usage, 
    validate_file_exists, 
    validate_text_content,
    get_file_size_mb,
    check_memory_available,
    cleanup_text
)

# Document processing imports
try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from striprtf.striprtf import rtf_to_text
    RTF_AVAILABLE = True
except ImportError:
    RTF_AVAILABLE = False


@dataclass
class LoadedDocument:
    """Structure representing a loaded document"""
    content: str
    source_file: str
    file_type: str
    metadata: Dict[str, Any]
    
    def __post_init__(self):
        """Clean content after initialization"""
        self.content = cleanup_text(self.content)


class LoaderError(Exception):
    """Document loading error"""
    pass


class DocumentLoader:
    """
    Multi-format document loader with memory management
    
    Supported formats:
    - .txt, .md (plain text)
    - .pdf (PyPDF2)
    - .docx (python-docx)
    - .rtf (striprtf)
    """
    
    SUPPORTED_EXTENSIONS = {'.txt', '.md', '.pdf', '.docx', '.rtf'}
    
    def __init__(self, max_file_size_mb: float = 50.0):
        self.max_file_size_mb = max_file_size_mb
    
    def load_document(self, file_path: str) -> LoadedDocument:
        """
        Load document from file path
        
        Args:
            file_path: Path to document file
            
        Returns:
            LoadedDocument with content and metadata
            
        Raises:
            LoaderError: If loading fails
        """
        log_memory_usage(f"Loading {file_path}")
        
        # Validate file
        if not validate_file_exists(file_path):
            raise LoaderError(f"File not found or not readable: {file_path}")
        
        # Check file size
        file_size_mb = get_file_size_mb(file_path)
        if file_size_mb and file_size_mb > self.max_file_size_mb:
            raise LoaderError(f"File too large: {file_size_mb:.1f}MB > {self.max_file_size_mb}MB")
        
        # Check available memory
        required_memory = (file_size_mb or 10) * 3  # Estimate 3x file size needed
        if not check_memory_available(required_memory):
            raise LoaderError(f"Insufficient memory for file size {file_size_mb:.1f}MB")
        
        # Determine file type and load
        file_path = Path(file_path)
        extension = file_path.suffix.lower()
        
        if extension not in self.SUPPORTED_EXTENSIONS:
            raise LoaderError(f"Unsupported file type: {extension}")
        
        try:
            # Load based on file type
            if extension in {'.txt', '.md'}:
                content, metadata = self._load_text(file_path)
            elif extension == '.pdf':
                content, metadata = self._load_pdf(file_path)
            elif extension == '.docx':
                content, metadata = self._load_docx(file_path)
            elif extension == '.rtf':
                content, metadata = self._load_rtf(file_path)
            else:
                raise LoaderError(f"Handler not implemented for: {extension}")
            
            # Validate loaded content
            if not validate_text_content(content):
                raise LoaderError(f"Invalid or empty content from: {file_path}")
            
            log_memory_usage(f"Loaded {file_path}")
            
            return LoadedDocument(
                content=content,
                source_file=str(file_path),
                file_type=extension[1:],  # Remove dot
                metadata=metadata
            )
            
        except Exception as e:
            raise LoaderError(f"Failed to load {file_path}: {e}")
    
    def _load_text(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Load plain text file (txt, md)"""
        # Try multiple encodings
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                
                metadata = {
                    'encoding': encoding,
                    'line_count': content.count('\n') + 1,
                    'char_count': len(content),
                    'file_size_mb': get_file_size_mb(str(file_path))
                }
                
                return content, metadata
                
            except UnicodeDecodeError:
                continue
        
        raise LoaderError(f"Could not decode text file with any encoding: {file_path}")
    
    def _load_pdf(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Load PDF file"""
        if not PDF_AVAILABLE:
            raise LoaderError("PyPDF2 not available. Install: pip install PyPDF2")
        
        content_parts = []
        page_count = 0
        
        try:
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                page_count = len(pdf_reader.pages)
                
                # Check if PDF is encrypted
                if pdf_reader.is_encrypted:
                    raise LoaderError("PDF is password protected")
                
                # Extract text from all pages
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            content_parts.append(page_text)
                    except Exception as e:
                        print(f"Warning: Could not extract text from page {page_num + 1}: {e}")
                
            content = "\n\n".join(content_parts)
            
            metadata = {
                'page_count': page_count,
                'pages_with_text': len(content_parts),
                'char_count': len(content),
                'file_size_mb': get_file_size_mb(str(file_path))
            }
            
            return content, metadata
            
        except Exception as e:
            raise LoaderError(f"PDF extraction failed: {e}")
    
    def _load_docx(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Load DOCX file"""
        if not DOCX_AVAILABLE:
            raise LoaderError("python-docx not available. Install: pip install python-docx")
        
        try:
            doc = DocxDocument(str(file_path))
            
            # Extract paragraphs
            paragraphs = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    paragraphs.append(text)
            
            content = "\n\n".join(paragraphs)
            
            metadata = {
                'paragraph_count': len(paragraphs),
                'total_paragraphs': len(doc.paragraphs),
                'char_count': len(content),
                'file_size_mb': get_file_size_mb(str(file_path))
            }
            
            return content, metadata
            
        except Exception as e:
            raise LoaderError(f"DOCX extraction failed: {e}")
    
    def _load_rtf(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Load RTF file"""
        if not RTF_AVAILABLE:
            raise LoaderError("striprtf not available. Install: pip install striprtf")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                rtf_content = f.read()
            
            # Convert RTF to plain text
            content = rtf_to_text(rtf_content)
            
            metadata = {
                'rtf_size': len(rtf_content),
                'text_size': len(content),
                'char_count': len(content),
                'file_size_mb': get_file_size_mb(str(file_path))
            }
            
            return content, metadata
            
        except Exception as e:
            raise LoaderError(f"RTF extraction failed: {e}")
    
    def get_supported_extensions(self) -> set[str]:
        """Get list of supported file extensions"""
        return self.SUPPORTED_EXTENSIONS.copy()
    
    def is_supported(self, file_path: str) -> bool:
        """Check if file type is supported"""
        extension = Path(file_path).suffix.lower()
        return extension in self.SUPPORTED_EXTENSIONS


# Convenience functions
def load_text(file_path: str, max_size_mb: float = 50.0) -> LoadedDocument:
    """Load single document - convenience function"""
    loader = DocumentLoader(max_file_size_mb=max_size_mb)
    return loader.load_document(file_path)